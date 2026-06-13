from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
SOURCE_LOGO = PUBLIC / "trustforce-logo.png"
RESOURCE_DIR = PUBLIC / "resources"
BUILD_DIR = ROOT / ".resource-build"

NAVY = RGBColor(8, 37, 72)
GOLD = RGBColor(196, 146, 35)
MUTED = RGBColor(86, 96, 110)
LIGHT_GOLD = "F7EED8"

RESOURCES = [
    (
        "conservatorship-checklist",
        "Conservatorship Checklist",
        "A practical checklist for organizing financial records, court documents, deadlines, and ongoing responsibilities.",
        "Conservators, family members, attorneys, and support professionals",
    ),
    (
        "guardianship-basics",
        "Guardianship Basics",
        "A plain-language overview of common guardian responsibilities, documentation needs, and care-related records.",
        "Guardians, families, caregivers, and referral partners",
    ),
    (
        "trustee-recordkeeping-guide",
        "Trustee Recordkeeping Guide",
        "A guide to organizing trust documents, beneficiary communications, distributions, receipts, and financial information.",
        "Trustees, beneficiaries, families, and professional advisors",
    ),
    (
        "veteran-fiduciary-support-guide",
        "Veteran Fiduciary Support Guide",
        "A future resource for veterans, families, and representatives managing benefit-related fiduciary responsibilities.",
        "Veterans, veteran families, fiduciaries, and advocates",
    ),
    (
        "court-accounting-preparation-guide",
        "Court Accounting Preparation Guide",
        "A structured guide for assembling organized records, summaries, and supporting documentation for review.",
        "Conservators, guardians, fiduciaries, attorneys, and families",
    ),
    (
        "fiduciary-glossary",
        "Fiduciary Glossary",
        "Simple definitions of common terms used in fiduciary, trust, conservatorship, guardianship, and estate matters.",
        "Clients, families, beneficiaries, fiduciaries, and professionals",
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run(run, size, color, bold=False, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_centered_text(doc, text, size, color, *, bold=False, italic=False, after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run(run, size, color, bold=bold, italic=italic)
    return paragraph


def build_favicon():
    image = Image.open(SOURCE_LOGO).convert("RGBA")
    mark = image.crop((410, 90, 1060, 610))
    canvas = Image.new("RGBA", (700, 700), (255, 255, 255, 255))
    mark.thumbnail((600, 600), Image.Resampling.LANCZOS)
    x = (canvas.width - mark.width) // 2
    y = (canvas.height - mark.height) // 2
    canvas.alpha_composite(mark, (x, y))
    canvas.resize((512, 512), Image.Resampling.LANCZOS).save(PUBLIC / "favicon.png")
    canvas.resize((180, 180), Image.Resampling.LANCZOS).save(PUBLIC / "apple-touch-icon.png")


def build_document(slug, title, description, audience):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("TRUSTFORCE ADVISORS  |  CLIENT RESOURCE CENTER")
    set_run(run, 8.5, MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run(
        "TheTrustForce.com  |  (629) 258-7878  |  Hello@TheTrustForce.com"
    )
    set_run(footer_run, 8.5, MUTED)

    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_after = Pt(4)
    logo_paragraph.add_run().add_picture(str(SOURCE_LOGO), width=Inches(3.35))

    add_centered_text(
        document,
        "GUIDE PREVIEW",
        10,
        GOLD,
        bold=True,
        after=10,
    )
    add_centered_text(document, title, 27, NAVY, bold=True, after=10)
    add_centered_text(
        document,
        description,
        12,
        MUTED,
        after=20,
    )

    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, LIGHT_GOLD)
    cell.vertical_alignment = 1
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_paragraph.paragraph_format.space_before = Pt(16)
    cell_paragraph.paragraph_format.space_after = Pt(4)
    status = cell_paragraph.add_run("COMING SOON")
    set_run(status, 18, GOLD, bold=True)
    message = cell.add_paragraph()
    message.alignment = WD_ALIGN_PARAGRAPH.CENTER
    message.paragraph_format.space_after = Pt(16)
    message_run = message.add_run(
        "This client resource is being developed and will be available in a future update."
    )
    set_run(message_run, 10.5, NAVY, bold=True)

    document.add_paragraph().paragraph_format.space_after = Pt(4)
    add_centered_text(document, "Designed For", 11, GOLD, bold=True, after=5)
    add_centered_text(document, audience, 11.5, NAVY, bold=True, after=18)
    add_centered_text(
        document,
        "Need assistance now?",
        12,
        NAVY,
        bold=True,
        after=5,
    )
    add_centered_text(
        document,
        "Schedule a consultation with TrustForce Advisors for professional, organized, and confidential support.",
        10.5,
        MUTED,
        after=10,
    )
    add_centered_text(
        document,
        "This preview is informational and does not provide legal advice.",
        8.5,
        MUTED,
        italic=True,
    )

    path = BUILD_DIR / f"{slug}.docx"
    document.save(path)
    return path


def main():
    PUBLIC.mkdir(exist_ok=True)
    RESOURCE_DIR.mkdir(exist_ok=True)
    BUILD_DIR.mkdir(exist_ok=True)
    build_favicon()
    for resource in RESOURCES:
        build_document(*resource)


if __name__ == "__main__":
    main()
